#!/usr/bin/env python3
"""Build the Bluebird FOI Insights capability paper as a branded .docx.

Run from the repo root:
    .venv\\Scripts\\python.exe docs\\build_capability_paper.py

Output: docs/Bluebird-FOI-Insights-capability-paper.docx

Visual identity is taken from https://axoquant.com (dark, engineering-led
brand): ink #0E1419, paper #EDEAE2, seal #2F9E6E, ember #C77B3A, with
Fraunces (display serif), Inter (body) and IBM Plex Mono (wordmark and
technical stamps). The web theme is dark; this template renders the same
palette on white for a printed document.

Sources: the design specs under docs/superpowers/specs/, the trend-window
decision (docs/decisions/2026-08-20-trend-window.md), README.md,
docs/deploy.md, and the verbatim corpus (data/corpus/data-notes.md).

Note on fonts: Fraunces, Inter and IBM Plex Mono are the site's fonts. Word
falls back gracefully (Georgia / Segoe UI / Consolas) on machines where they
are not installed.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Bluebird-FOI-Insights-capability-paper.docx"

# Brand palette (axoquant.com)
INK = RGBColor(0x0E, 0x14, 0x19)
SEAL = RGBColor(0x2F, 0x9E, 0x6E)
SEAL_DEEP = RGBColor(0x1F, 0x6B, 0x4A)
EMBER = RGBColor(0xC7, 0x7B, 0x3A)
PAPER = RGBColor(0xED, 0xEA, 0xE2)
PAPER_DIM = RGBColor(0xB7, 0xB4, 0xAA)
SLATE = RGBColor(0x6E, 0x76, 0x80)

HEX = {
    "ink": "0E1419",
    "seal": "2F9E6E",
    "seal_deep": "1F6B4A",
    "ember": "C77B3A",
    "paper": "EDEAE2",
    "paper_dim": "B7B4AA",
    "white": "FFFFFF",
}

DISPLAY = "Fraunces"
SANS = "Inter"
MONO = "IBM Plex Mono"


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def _set_run(run, font=SANS, size=11, color=INK, bold=False, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def shade_paragraph(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def edge_rule(p, edge, color, sz=8):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "8")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def set_table_borders(table, color, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def add_page_field(p):
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    return run


# ---------------------------------------------------------------------------
# Content builders
# ---------------------------------------------------------------------------

def body(doc, text, size=11, italic=False, color=INK, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    _set_run(p.add_run(text), SANS, size, color, bold, italic)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    _set_run(p.add_run(text), SANS, 11, INK)
    return p


def arrow_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.18)
    _set_run(p.add_run("→  "), MONO, 11, EMBER, bold=True)
    _set_run(p.add_run(text), SANS, 11, INK)
    return p


def h1(doc, num, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(f"{num:02d}   "), MONO, 13, EMBER, bold=True)
    _set_run(p.add_run(title), DISPLAY, 17, INK, bold=True)
    return p


def h1_plain(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(title), DISPLAY, 17, INK, bold=True)
    return p


def h2(doc, title):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(title), DISPLAY, 13, SEAL_DEEP, bold=True)
    return p


def callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, HEX["paper"])
    edge_rule(p, "left", HEX["seal"], sz=24)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    _set_run(p.add_run(label), MONO, 9, SEAL_DEEP, bold=True)
    q = doc.add_paragraph()
    q.paragraph_format.space_after = Pt(8)
    q.paragraph_format.line_spacing = 1.15
    shade_paragraph(q, HEX["paper"])
    edge_rule(q, "left", HEX["seal"], sz=24)
    q.paragraph_format.left_indent = Inches(0.15)
    q.paragraph_format.right_indent = Inches(0.15)
    _set_run(q.add_run(text), SANS, 11.5, INK)


def build_stamp(doc, text, color=SLATE):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _set_run(p.add_run(text), MONO, 8.5, color)
    return p


def branded_table(doc, header, rows, widths):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    set_table_borders(table, HEX["paper_dim"], sz=4)
    for i, h in enumerate(header):
        c = table.rows[0].cells[i]
        shade_cell(c, HEX["seal_deep"])
        c.width = Inches(widths[i])
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        _set_run(p.add_run(h), SANS, 10, PAPER, bold=True)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        fill = HEX["paper"] if ri % 2 == 1 else HEX["white"]
        for i, val in enumerate(row):
            c = cells[i]
            shade_cell(c, fill)
            c.width = Inches(widths[i])
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _set_run(p.add_run(val), SANS, 9.5, INK)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return table


def contents(doc):
    h1_plain(doc, "Contents")
    items = [
        ("1.  Data issues found in the current solution", False),
        ("2.  Options considered, and the decision we took", False),
        ("3.  Analysis of the Power BI dashboard and the source data", False),
        ("4.  What we built, and how lineage provides certainty", False),
        ("5.  The new solution at a glance, for a new user", False),
        ("6.  What the capability makes possible", False),
        ("6a.  More insight for the public from data that already exists", True),
        ("6b.  A chat interface across the data for general users", True),
        ("6c.  Risk-based insight for internal users, with forecasting and "
         "classification", True),
        ("6d.  Linking the FOI data with an agency's own data", True),
    ]
    for text, indent in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Inches(0.25)
        _set_run(p.add_run(text), SANS, 11, INK)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# The paper
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    # A4, branded margins
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.1)
    sec.footer_distance = Cm(1.1)
    sec.different_first_page_header_footer = True

    # Blank first-page header so the cover carries no running header. The
    # default header below only appears from page 2 onward.
    first_hdr = sec.first_page_header
    first_hdr.paragraphs[0].text = ""

    # Running header (blank on cover)
    hdr = sec.header.paragraphs[0]
    hdr.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _set_run(hdr.add_run("BLUEBIRD FOI INSIGHTS  ·  FOI STATISTICS CAPABILITY PAPER"),
             MONO, 8, SLATE)
    _set_run(hdr.add_run("\tAXO_QUANT_"), MONO, 8, SEAL)
    edge_rule(hdr, "bottom", HEX["seal"], sz=8)

    # Running footer (blank on cover)
    ftr = sec.footer.paragraphs[0]
    ftr.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _set_run(ftr.add_run("GOVERNED SOVEREIGN AI  ·  AXO_QUANT_"), MONO, 8, SLATE)
    _set_run(ftr.add_run("\tPAGE "), MONO, 8, SLATE)
    fld = add_page_field(ftr)
    _set_run(fld, MONO, 8, SLATE)
    edge_rule(ftr, "top", HEX["paper_dim"], sz=4)

    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(11)

    doc.core_properties.title = "Bluebird FOI Insights — Capability Paper"
    doc.core_properties.subject = "FOI statistics capability paper"
    doc.core_properties.author = "Bluebird Advisory"

    # ------------------------------------------------------------------
    # Cover block
    # ------------------------------------------------------------------
    mast = doc.add_paragraph()
    _set_run(mast.add_run("AXO_QUANT_  /  BLUEBIRD FOI INSIGHTS"), MONO, 9, SLATE)
    edge_rule(mast, "bottom", HEX["seal"], sz=12)

    doc.add_paragraph()
    t = doc.add_paragraph()
    _set_run(t.add_run("Bluebird FOI Insights"), DISPLAY, 30, INK, bold=True)
    t.paragraph_format.space_after = Pt(4)

    s = doc.add_paragraph()
    _set_run(s.add_run("From a static dashboard to a lineage-traced FOI statistics "
                       "platform"), SANS, 15, SLATE, italic=True)
    s.paragraph_format.space_after = Pt(6)

    m = doc.add_paragraph()
    _set_run(m.add_run("Internal capability paper   ·   24 August 2026   ·   "
                       "Bluebird Advisory"), MONO, 10, SLATE)
    m.paragraph_format.space_after = Pt(14)

    callout(
        doc,
        "CONTEXT",
        "Bluebird FOI Insights is a hosted platform built on the AxoQuant sovereign "
        "stack. It reproduces the published Australian Government FOI statistics "
        "pages as twelve interactive views, adds a natural-language ask, chat and "
        "report path, and records lineage for every figure so each number can be "
        "traced to a pinned source file and a recorded calculation. This paper "
        "records the data issues we found in the current Power BI solution, the "
        "decision we took on presenting the trend, what our analysis taught us, "
        "what we built, a primer for a new user, and the capability the result "
        "opens up.",
    )
    build_stamp(doc, "BUILD: 2026-08-24   ·   DATA: data.gov.au FOI statistics "
                     "(pinned snapshot)   ·   BASIS: published figures   ·   "
                     "LINEAGE: verified")

    contents(doc)

    # ------------------------------------------------------------------
    # 1. Data issues
    # ------------------------------------------------------------------
    h1(doc, 1, "Data issues found in the current solution")
    body(doc, "The data behind the current solution is sound at the point of "
              "publication: agencies report figures that the publisher "
              "quality-checks. The problems are structural, in the way the "
              "published files are shaped and in what a naive reader assumes "
              "about them. We catalogued seven concrete issues before building "
              "anything.")
    branded_table(
        doc,
        ["Issue", "What happens", "Why it matters"],
        [
            [
                "Cumulative versus single-quarter",
                "The current file (2025-26 Q1 to Q3) is cumulative: 34,418 "
                "requests received. The Power BI report shows the single-quarter "
                "Q1 figure, 12,359. Summing the cumulative column returns 34,418, "
                "not 12,359.",
                "Any tool that reads the file naively shows the wrong headline "
                "number, and the report's single-quarter view cannot be reproduced "
                "from the file alone.",
            ],
            [
                "Total rows do not re-sum",
                "The sheet's own Total (34,810) differs from the sum of the "
                "agency rows (34,418) because of an extra transfer line.",
                "Recomputing a total from the rows produces a different, wrong "
                "number. The published Total must be treated as the trusted value.",
            ],
            [
                "Note rows mixed into the data",
                "Rows prefixed with an x, for example x Norfolk Island (external "
                "territory), annotate the table but parse as agency rows.",
                "A naive parse treats notes as agencies and adds phantom rows to "
                "every aggregation.",
            ],
            [
                "Machinery of Government renames",
                "Agencies change name or merge across years, for example Human "
                "Services to Social Services. The notes state that renamed "
                "agencies appear under their most recent name, including for "
                "periods where the former name applied.",
                "Year-over-year comparisons break unless names are resolved once "
                "against a curated map.",
            ],
            [
                "The trend axis disagrees with the page text",
                "The report's trend views display October 2023 to September 2025, "
                "eight quarters or about two financial years. The page text "
                "describes a five-year view.",
                "A replicating tool must decide which claim to honour, and the "
                "decision needs recording.",
            ],
            [
                "Per-quarter data stops in 2018-19",
                "Quarterly files exist only through 2018-19. The annual files from "
                "2019-20 onward are financial-year totals only.",
                "Individual quarters for recent years cannot be sourced. "
                "Reconstructing them would mean inventing data.",
            ],
            [
                "Decision and timeliness measures present but not read",
                "The workbooks publish decisions, decision outcomes and timeliness "
                "on dedicated sheets. An early read of the data pulled only the "
                "Request numbers sheet.",
                "Six chart pages had no data to show even though the source "
                "workbook contained the figures.",
            ],
        ],
        widths=[1.35, 2.65, 2.4],
    )
    body(doc, "The common thread is that every issue sits at the boundary between "
              "the published files and the way a reader interprets them. Each is "
              "small; together they make a naive reading unreliable. The last one "
              "changed the shape of the work. The data was present but unused, "
              "which pointed to building a proper ingest rather than another "
              "dashboard over the raw files.")

    # ------------------------------------------------------------------
    # 2. Options and decision
    # ------------------------------------------------------------------
    h1(doc, 2, "Options considered, and the decision we took")
    body(doc, "Two questions needed answers before anything could be built: which "
              "trend window to show, and how to present single-quarter figures "
              "that the current file does not directly contain. Three options were "
              "on the table.")
    branded_table(
        doc,
        ["Option", "Approach", "Outcome"],
        [
            [
                "1. Reconstruct per-quarter figures",
                "Derive individual quarters for 2019-20 onward by differencing or "
                "synthesising from the annual totals.",
                "Rejected. The published sources do not contain quarterly detail "
                "for those years. Any reconstruction is invented data, and invented "
                "data in a government-facing demo is a credibility risk.",
            ],
            [
                "2. Published-data version of the claim",
                "Take the published single-quarter Q1 2025-26 figures as ground "
                "truth for the headline numbers, and the published financial-year "
                "totals (2019-20 through 2024-25, plus the Q1 to Q3 cumulative "
                "file) for the trend views. No quarter is reconstructed.",
                "Chosen. Honest to the published files, consistent with the "
                "headline figures, and every basis is labelled.",
            ],
            [
                "3. Copy the report's trend view exactly",
                "Replicate the eight-quarter October 2023 to September 2025 view "
                "the report actually displays.",
                "Not chosen. It would copy the discrepancy rather than resolve it, "
                "and would contradict the page's five-year claim.",
            ],
        ],
        widths=[1.5, 3.0, 1.9],
    )
    body(doc, "The decision is recorded in docs/decisions/2026-08-20-trend-window.md, "
              "agreed 2026-08-20. We chose a single-quarter Q1 2025-26 headline and "
              "a financial-year five-year trend from the published annual files, "
              "with no per-quarter reconstruction. Where the report's trend view "
              "disagreed with its own page text, the platform follows the "
              "published-data version of the claim, and the lineage notes the "
              "discrepancy wherever it is relevant.")
    body(doc, "The decision was recorded in writing because a government audience "
              "would catch a wrong trend window, and because the appealing option "
              "was the easy one. Reproducing the report exactly is straightforward "
              "until someone asks where a quarter came from. Option 2 is the one "
              "that can be answered from the record.")

    # ------------------------------------------------------------------
    # 3. Analysis of the Power BI dashboard
    # ------------------------------------------------------------------
    h1(doc, 3, "Analysis of the Power BI dashboard and the source data")

    h2(doc, "What was good")
    bullet(doc, "The report is a complete reference product. Twelve pages cover "
                "volume, outcomes, timeliness and top contributors with filters, "
                "and they define the information surface the replacement needed "
                "to match.")
    bullet(doc, "The published headline figures are usable as ground truth. The "
                "golden numbers, 12,359 requests received in Q1 2025-26 and the "
                "rest, anchor the whole build, and they hold up under verification.")
    bullet(doc, "The data notes and disclaimer are a definitional authority. They "
                "explain the renames, the quarterly-versus-financial-year basis, "
                "and what personal information means in this dataset. Having that "
                "text verbatim made every subsequent judgement defensible.")

    h2(doc, "What worked in the analysis")
    bullet(doc, "Every quirk was discoverable from the files and the notes "
                "together. Nothing required guessing; each oddity had a documented "
                "cause, whether a transfer line, a note row, or a rename.")
    bullet(doc, "The workbooks were structurally consistent across years, which "
                "made a header-driven parse viable for the decision and timeliness "
                "sheets. That consistency is what let us close the gap described "
                "in section 1 without hand-editing years.")

    h2(doc, "What did not work")
    bullet(doc, "There is no lineage. A figure in the Power BI report cannot be "
                "traced to a source cell, so a wrong-looking number cannot be "
                "audited and a correct one cannot be defended. For a "
                "government-facing product this is the largest gap.")
    bullet(doc, "The report's trend view contradicts its own page text: eight "
                "quarters displayed against a five-year claim.")
    bullet(doc, "The current file is cumulative, so the single-quarter headline "
                "the report shows cannot be reproduced from the file alone.")
    bullet(doc, "A user cannot safely re-sort or re-sum the data. Because the "
                "Total row does not re-sum, any recomputation yields a different "
                "number.")
    bullet(doc, "The reports are static exports. There is no programmatic access "
                "to the figures, no way to ask a question of the data, and no way "
                "to carry the analysis beyond what the page already shows.")

    body(doc, "The short version: the dashboard is a good display of the data and "
              "a weak source of truth. Its value is the information surface; its "
              "weakness is that nothing behind it can be verified or extended.")

    # ------------------------------------------------------------------
    # 4. What we built, and how lineage provides certainty
    # ------------------------------------------------------------------
    h1(doc, 4, "What we built, and how lineage provides certainty")
    body(doc, "The response was to build a normalising ingest and a lineage-traced "
              "platform rather than another dashboard over the raw files. Six "
              "pieces of machinery deliver the certainty both the users and the "
              "OAIC need.")
    bullet(doc, "A normalising ingest. The ingest resolves every quirk once: MoG "
                "renames from a curated map, note rows, trusted totals, and the "
                "cumulative-versus-single-quarter discrepancy. It writes canonical "
                "long-form facts, one row per agency, period, measure, bucket and "
                "value, with the basis recorded on the row.")
    bullet(doc, "A durable store. Facts land in Postgres: foi_datasets pins each "
                "source snapshot (path, URL, sha256, size, download time), and "
                "foi_facts holds the immutable rows. Refresh inserts a new "
                "snapshot; it never updates an old one, so older views keep their "
                "exact lineage.")
    bullet(doc, "A golden boot check. At startup the app asserts the published "
                "Q1 2025-26 headline figures against the loaded data and refuses "
                "to boot on a mismatch. The service cannot serve a wrong headline "
                "number without failing loudly first.")
    bullet(doc, "A lineage ledger. A hybrid ledger, a JSONL event stream plus "
                "Postgres tables, records for every figure and every request the "
                "dataset snapshot, the calculation, the rows it used and a hash "
                "of those rows, and the outcome. Replay verification recomputes "
                "each recorded operation rather than trusting the stored value.")
    bullet(doc, "The model cannot write a digit. The chat and report paths "
                "produce structure and keys; the platform computes every number. "
                "Answers carry citation pointers that resolve against the "
                "recorded transcript, and an unresolvable pointer fails rather "
                "than printing a guess.")
    bullet(doc, "A basis on every figure. The renderer prints whether a figure is "
                "single-quarter, cumulative or financial-year beside it, and the "
                "lineage page shows the same.")
    body(doc, "For the users, certainty is structural. Any figure on the site can "
              "be traced to a pinned source file, a recorded calculation and a "
              "verifiable hash. For the OAIC, the same machinery answers the "
              "question of where a number came from. The answer is a path you can "
              "walk.")

    # ------------------------------------------------------------------
    # 5. Primer for a new user
    # ------------------------------------------------------------------
    h1(doc, 5, "The new solution at a glance, for a new user")
    bullet(doc, "What it is. A hosted site that reproduces the twelve FOI "
                "statistics pages with interactive charts, adds a natural-language "
                "ask, chat and report path, and records lineage for everything. "
                "The statistics pages are open to the public; a login-gated Chat "
                "and reports section holds the interactive Q&A.")
    bullet(doc, "Data. The pinned data.gov.au FOI statistics snapshot, normalised "
                "into canonical facts. Nothing is fetched live when a page loads; "
                "the snapshot is baked into the deployment.")
    bullet(doc, "The rule. The platform never invents a number. Figures come from "
                "the frame or the verbatim corpus; the model supplies language, "
                "not digits.")
    bullet(doc, "The pages. Twelve pages, from at a glance through requests, "
                "decisions, outcomes, timeliness and reference, rendered from "
                "platform-computed figures with basis labels and interactive "
                "charts.")
    bullet(doc, "Asking questions. The gated chat answers grounded questions over "
                "the data and corpus with citations. The report turns a request "
                "into a real figure with its lineage. Out-of-scope requests are "
                "refused cleanly and point to an escalation email.")
    bullet(doc, "Governance. Defence in depth: a deterministic scope screen before "
                "the model, a prompt-level scope block, a jailbreak scan, a "
                "read-only tool sandbox, and an identity stovepipe that discloses "
                "only that the model is powered by the sovereign stack.")
    bullet(doc, "Running it. python scripts/serve.py starts a local copy on port "
                "8095. Deployment is python scripts/deploy.py with dry-run, check "
                "and no-restart modes; the full runbook is in docs/deploy.md.")
    bullet(doc, "Where to look. The lineage page for any artifact, the API access "
                "page, and the data notes page, which carries the publisher's "
                "notes verbatim.")

    # ------------------------------------------------------------------
    # 6. Capability
    # ------------------------------------------------------------------
    h1(doc, 6, "What the capability makes possible")
    body(doc, "Everything below runs on the same normalised, lineage-traced facts. "
              "The capability is the data spine plus the discipline around it; "
              "each of these is an option that becomes cheap once the spine "
              "exists.")

    h2(doc, "6a. More insight for the public from data that already exists")
    bullet(doc, "Interactive charts with hover values, zoom, legend toggles and "
                "live filters by portfolio, agency, type and period give the "
                "public a richer read of the published statistics than static "
                "images.")
    bullet(doc, "A read-only API exposes the figures programmatically, so analysts "
                "and researchers can pull the numbers rather than scraping a "
                "report.")
    bullet(doc, "Every figure carries its basis and its lineage, so a member of "
                "the public can verify what they are reading.")
    bullet(doc, "Examples already answerable today: which agencies contribute most "
                "to requests received, and how timeliness has changed across "
                "agencies.")

    h2(doc, "6b. A chat interface across the data for general users")
    body(doc, "A grounded chat answers questions from the corpus and the catalog, "
              "with citations under each answer and a deterministic fallback so "
              "the chat never dies and never fabricates. Example answers:")
    arrow_item(doc, "How many requests were received in Q1 2025-26? Returns 12,359 "
                    "with a citation.")
    arrow_item(doc, "Which agencies increased their refusal rate most between "
                    "FY23 and FY24? A compare-period answer over the canonical "
                    "facts.")
    arrow_item(doc, "Where do requests concentrate by agency? Grounded in the "
                    "notes and the data.")
    arrow_item(doc, "Top agencies by requests received in Q1 2025-26. A report "
                    "card with the real figure and its lineage.")
    body(doc, "Requests outside scope are refused cleanly and point to an "
              "escalation path, so the interface stays a statistics tool rather "
              "than becoming a general assistant.")

    h2(doc, "6c. Risk-based insight for internal users, with forecasting and "
            "classification")
    body(doc, "The same canonical facts can support internal risk views for an FOI "
              "agency. These are read-only analyses over published data, each "
              "computed by the platform and labelled with its basis.")
    arrow_item(doc, "Timeliness risk: agencies whose within-statutory performance "
                    "is degrading, read from the published timeliness series.")
    arrow_item(doc, "Outcome movers: agencies whose mix of granted, refused and "
                    "withdrawn decisions is shifting.")
    arrow_item(doc, "Time-series forecasting: projections of request volume and "
                    "timeliness from the financial-year series, computed by the "
                    "platform, never by the model.")
    arrow_item(doc, "Tabular classification: agency risk profiles built from "
                    "volume, timeliness and outcome features, giving a repeatable "
                    "and explainable partition of agencies.")
    body(doc, "There are honest limits. Some measures exist only at single-quarter "
              "or financial-year granularity, so risk views are built on what is "
              "published and clearly labelled where they derive from it. The value "
              "is a repeatable, explainable risk picture rather than a claim of "
              "prediction.")

    h2(doc, "6d. Linking the FOI data with an agency's own data")
    body(doc, "The normalised facts are a clean spine that can join to other data, "
              "including the datasets an agency already holds in its own datalake. "
              "Two consequences follow.")
    bullet(doc, "Published statistics can be blended with internal operational "
                "data, for example case workloads, resourcing or contact data, to "
                "inform demand and resourcing forecasts.")
    bullet(doc, "The lineage machinery makes the join auditable: every blended "
                "figure traces to its sources and the join rule, so a derived "
                "number is as defensible as a published one. The same figure and "
                "DSL layer can be pointed at the joined dataset rather than "
                "rebuilt.")
    body(doc, "This is offered as capability. The data-sharing and access "
              "arrangements would need to be settled before anything is built, and "
              "the platform's lineage discipline is exactly what makes such a join "
              "reviewable.")

    # ------------------------------------------------------------------
    # Close
    # ------------------------------------------------------------------
    doc.add_paragraph()
    body(doc, "The published FOI data has real structural quirks, and the existing "
              "dashboard displays them without lineage. The replacement turns the "
              "same published data into a verified, queryable platform: every "
              "figure traces to a source, every request is answerable or cleanly "
              "refused, and the path from data to insight is a record you can "
              "follow.", space_after=14)
    build_stamp(doc, "AXO_QUANT_BUILT  ·  BLUEBIRD FOI INSIGHTS  ·  "
                     "LINEAGE-TRACED  ·  SOURCES: DATA.GOV.AU", SEAL)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
